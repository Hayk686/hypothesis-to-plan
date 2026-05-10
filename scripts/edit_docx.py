#!/usr/bin/env python3
"""DOCX editor: read, append, find/replace, create.

Emits a single JSON line to stdout on the last line, matching the
convention used by sibling wrappers (status/files/count, plus optional
content/error fields).
"""

import argparse
import json
import os
import sys

try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")
except AttributeError:
    pass

try:
    from docx import Document
except ImportError:
    print(json.dumps({
        "status": "error",
        "error": "python-docx is not installed. Run: pip install python-docx",
        "files": [],
        "count": 0,
    }, ensure_ascii=False))
    sys.exit(2)

OUTPUT_DIR = os.path.join("output", "documents")


def emit(payload: dict, code: int = 0) -> None:
    print(json.dumps(payload, ensure_ascii=False))
    sys.exit(code)


def strip_file_wrapper(path: str) -> str:
    p = path.strip()
    if p.startswith("[file:") and p.endswith("]"):
        p = p[6:-1]
    return p


def resolve_output(input_path: str, explicit: str | None) -> str:
    if explicit:
        out = explicit
        os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)
        return out
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    base = os.path.splitext(os.path.basename(input_path))[0]
    return os.path.join(OUTPUT_DIR, f"{base}.docx")


def cmd_read(args) -> None:
    src = strip_file_wrapper(args.input)
    if not os.path.isfile(src):
        emit({"status": "error", "error": f"file not found: {src}", "files": [], "count": 0}, 1)
    doc = Document(src)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text:
                        parts.append(p.text)
    emit({"status": "ok", "content": "\n".join(parts), "files": [], "count": 0})


def cmd_append(args) -> None:
    src = strip_file_wrapper(args.input)
    if not os.path.isfile(src):
        emit({"status": "error", "error": f"file not found: {src}", "files": [], "count": 0}, 1)
    doc = Document(src)
    text = args.text.replace("\\n", "\n")
    for line in text.split("\n"):
        doc.add_paragraph(line)
    out = resolve_output(src, args.output)
    doc.save(out)
    emit({"status": "ok", "files": [os.path.abspath(out)], "count": 1})


def cmd_replace(args) -> None:
    src = strip_file_wrapper(args.input)
    if not os.path.isfile(src):
        emit({"status": "error", "error": f"file not found: {src}", "files": [], "count": 0}, 1)
    doc = Document(src)
    needle = args.find
    repl = args.replace
    hits = 0

    def replace_in_paragraph(p) -> int:
        if needle not in p.text:
            return 0
        new_text = p.text.replace(needle, repl)
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = new_text
        else:
            p.add_run(new_text)
        return p.text.count(repl) if repl else 1

    for p in doc.paragraphs:
        if needle in p.text:
            hits += p.text.count(needle)
            replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if needle in p.text:
                        hits += p.text.count(needle)
                        replace_in_paragraph(p)

    out = resolve_output(src, args.output)
    doc.save(out)
    emit({"status": "ok", "files": [os.path.abspath(out)], "count": 1, "replacements": hits})


def cmd_create(args) -> None:
    out = args.output
    if not out:
        emit({"status": "error", "error": "--output is required for create", "files": [], "count": 0}, 1)
    os.makedirs(os.path.dirname(os.path.abspath(out)) or ".", exist_ok=True)
    doc = Document()
    text = (args.text or "").replace("\\n", "\n")
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(out)
    emit({"status": "ok", "files": [os.path.abspath(out)], "count": 1})


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Edit .docx files: read/append/replace/create")
    sub = parser.add_subparsers(dest="action", required=True)

    p_read = sub.add_parser("read", help="Extract plain text from a .docx")
    p_read.add_argument("--input", required=True)
    p_read.set_defaults(func=cmd_read)

    p_app = sub.add_parser("append", help="Append text as new paragraph(s) to the end")
    p_app.add_argument("--input", required=True)
    p_app.add_argument("--text", required=True)
    p_app.add_argument("--output", help="Optional output path. Defaults to output\\documents\\<name>.docx")
    p_app.set_defaults(func=cmd_append)

    p_rep = sub.add_parser("replace", help="Find/replace text inside a .docx")
    p_rep.add_argument("--input", required=True)
    p_rep.add_argument("--find", required=True)
    p_rep.add_argument("--replace", required=True)
    p_rep.add_argument("--output", help="Optional output path. Defaults to output\\documents\\<name>.docx")
    p_rep.set_defaults(func=cmd_replace)

    p_new = sub.add_parser("create", help="Create a new .docx from text")
    p_new.add_argument("--output", required=True)
    p_new.add_argument("--text", default="")
    p_new.set_defaults(func=cmd_create)

    return parser


def main() -> None:
    args = build_parser().parse_args()
    try:
        args.func(args)
    except Exception as e:
        emit({"status": "error", "error": str(e), "files": [], "count": 0}, 1)


if __name__ == "__main__":
    main()
